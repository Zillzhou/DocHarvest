"""
文档格式转换模块
支持将飞书文档内容转换为Markdown、Word (docx) 和 PDF格式
"""
import os
import re
import logging
import platform
from typing import Dict, Any, List, Optional

# Word导出依赖
try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
except ImportError:
    Document = None

# PDF导出依赖
try:
    import markdown
    from xhtml2pdf import pisa
except ImportError:
    markdown = None
    pisa = None


class DocumentConverter:
    """多格式文档转换器"""
    
    def __init__(self):
        self.logger = logging.getLogger(__name__)
    
    def to_markdown(self, doc_content: Dict[str, Any], doc_metadata: Dict[str, Any] = None) -> str:
        """
        转换为Markdown格式
        """
        if not doc_content:
            self.logger.error("文档内容为空")
            return ""
        
        markdown_lines = []
        
        # 添加文档标题
        if doc_metadata and doc_metadata.get("title"):
            title = doc_metadata.get("title", "未命名文档")
            markdown_lines.append(f"# {title}\n")
        
        # 获取纯文本内容
        content = doc_content.get("content", "")
        
        if content:
            markdown_lines.append(self._format_content(content))
        else:
            self.logger.warning("文档内容为空")
            markdown_lines.append("*文档内容为空*")
        
        return "\n".join(markdown_lines)
    
    def to_docx(self, doc_content: Dict[str, Any], doc_metadata: Dict[str, Any] = None, save_path: str = None) -> bool:
        """
        转换为Word格式
        """
        if Document is None:
            self.logger.error("未安装python-docx库，无法导出Word")
            return False
            
        if not doc_content:
            return False
            
        try:
            doc = Document()
            
            # 添加标题
            if doc_metadata and doc_metadata.get("title"):
                title = doc_metadata.get("title", "未命名文档")
                doc.add_heading(title, 0)
            
            # 获取内容
            content = doc_content.get("content", "")
            if content:
                # 处理内容
                lines = content.split('\n')
                prev_empty = False
                
                for line in lines:
                    stripped = line.strip()
                    
                    if not stripped:
                        if not prev_empty:
                            doc.add_paragraph("")
                            prev_empty = True
                        continue
                    
                    prev_empty = False
                    
                    # 简单的标题检测
                    if self._is_likely_heading(stripped, line):
                        level = self._detect_heading_level(stripped)
                        # Word标题级别从1开始，且不能超过9
                        level = min(max(level, 1), 9)
                        doc.add_heading(stripped, level=level)
                    else:
                        doc.add_paragraph(stripped)
            
            if save_path:
                doc.save(save_path)
                return True
            return False
            
        except Exception as e:
            self.logger.error(f"导出Word失败: {str(e)}")
            return False

    def to_pdf(self, doc_content: Dict[str, Any], doc_metadata: Dict[str, Any] = None, save_path: str = None) -> bool:
        """
        转换为PDF格式
        """
        if markdown is None or pisa is None:
            self.logger.error("未安装markdown或xhtml2pdf库，无法导出PDF")
            return False
            
        try:
            # 先转换为Markdown
            md_content = self.to_markdown(doc_content, doc_metadata)
            
            # Markdown转HTML
            html_content = markdown.markdown(md_content)
            
            # 添加基本的HTML结构和CSS
            # 注意：xhtml2pdf对中文字体支持有限，这里使用基本样式
            # 中文可能显示为方块，这是xhtml2pdf的已知限制
            full_html = f"""
            <html>
            <head>
                <meta charset="utf-8">
                <style>
                    body {{
                        font-family: sans-serif;
                        font-size: 12pt;
                        line-height: 1.6;
                    }}
                    h1, h2, h3, h4, h5, h6 {{ 
                        margin-top: 20px; 
                        margin-bottom: 10px; 
                        font-weight: bold;
                    }}
                    h1 {{ font-size: 24pt; }}
                    h2 {{ font-size: 20pt; }}
                    h3 {{ font-size: 16pt; }}
                    p {{ margin-bottom: 10px; line-height: 1.5; }}
                    pre {{ background-color: #f5f5f5; padding: 10px; border-radius: 5px; }}
                    code {{ background-color: #f5f5f5; padding: 2px 4px; border-radius: 3px; }}
                </style>
            </head>
            <body>
                {html_content}
            </body>
            </html>
            """
            
            if save_path:
                with open(save_path, "wb") as f:
                    pisa_status = pisa.CreatePDF(full_html, dest=f, encoding='utf-8')

                
                return not pisa_status.err
            return False
            
        except Exception as e:

            self.logger.error(f"导出PDF失败: {str(e)}")
            return False

    def _get_chinese_font_path(self) -> Optional[str]:
        """获取中文字体路径"""
        system = platform.system()
        if system == "Windows":
            # 尝试常见的Windows中文字体
            fonts = [
                "C:\\Windows\\Fonts\\msyh.ttc", # 微软雅黑
                "C:\\Windows\\Fonts\\simhei.ttf", # 黑体
                "C:\\Windows\\Fonts\\simsun.ttc", # 宋体
            ]
            for font in fonts:
                if os.path.exists(font):
                    return font.replace("\\", "/") # xhtml2pdf需要正斜杠或转义
        elif system == "Darwin": # macOS
            fonts = [
                "/System/Library/Fonts/PingFang.ttc",
                "/Library/Fonts/Arial Unicode.ttf"
            ]
            for font in fonts:
                if os.path.exists(font):
                    return font
        # Linux等其他系统暂不处理，或者需要用户指定
        return None

    def _format_content(self, content: str) -> str:
        """格式化内容 (复用原逻辑)"""
        lines = content.split('\n')
        formatted_lines = []
        
        prev_empty = False
        for line in lines:
            stripped = line.strip()
            
            if not stripped:
                if not prev_empty:
                    formatted_lines.append("")
                    prev_empty = True
                continue
            
            prev_empty = False
            
            if self._is_likely_heading(stripped, line):
                level = self._detect_heading_level(stripped)
                formatted_lines.append(f"{'#' * level} {stripped}")
            else:
                formatted_lines.append(line)
        
        return "\n".join(formatted_lines)
    
    def _is_likely_heading(self, stripped: str, original: str) -> bool:
        """判断是否为标题 (复用原逻辑)"""
        if len(stripped) > 80:
            return False
        
        if re.match(r'^[\d一二三四五六七八九十]+[、\.\s]', stripped):
            return True
        
        if len(stripped) < 30 and not stripped.endswith(('。', '，', ',', '.', '；', ';')):
            return True
        
        return False
    
    def _detect_heading_level(self, text: str) -> int:
        """检测标题级别 (复用原逻辑)"""
        if re.match(r'^一[、\.]', text) or re.match(r'^1[、\.]', text):
            return 2
        elif re.match(r'^[\d一二三四五六七八九十]+[、\.]', text):
            return 3
        elif len(text) < 20:
            return 2
        else:
            return 3
